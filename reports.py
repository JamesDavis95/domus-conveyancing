from io import BytesIO
from typing import Any, Dict, List
from docx import Document
from docx.shared import Pt

def build_docx(matter: Dict[str, Any],
               risks: List[Dict[str, str]],
               findings: List[Dict[str, str]]) -> bytes:
    doc = Document()
    title = doc.add_heading("Conveyancing Report", level=0)
    for r in title.runs:
        r.font.size = Pt(16)

    doc.add_paragraph(f"Reference: {matter.get('ref','')}")
    doc.add_paragraph(f"Created:  {matter.get('created_at','')}")

    doc.add_heading("Risks", level=1)
    if risks:
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells
        hdr[0].text = "Code"
        hdr[1].text = "Level"
        for r in risks:
            row = t.add_row().cells
            row[0].text = str(r.get("code",""))
            row[1].text = str(r.get("level",""))
    else:
        doc.add_paragraph("None.")

    doc.add_heading("Findings", level=1)
    if findings:
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells
        hdr[0].text = "Kind"
        hdr[1].text = "Value"
        for f in findings:
            row = t.add_row().cells
            row[0].text = str(f.get("kind",""))
            row[1].text = str(f.get("value",""))
    else:
        doc.add_paragraph("None.")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def build_json(matter: Dict[str, Any],
               risks: List[Dict[str,str]],
               findings: List[Dict[str,str]]) -> Dict[str, Any]:
    return {
        "matter": {
            "id": matter.get("id"),
            "ref": matter.get("ref"),
            "created_at": matter.get("created_at"),
        },
        "risks": risks or [],
        "findings": findings or [],
    }
