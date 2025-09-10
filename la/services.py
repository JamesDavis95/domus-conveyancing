import os, json, hashlib
from typing import Tuple, List
from sqlalchemy.orm import Session
from docx import Document
from la.models import LAMatter, LAOrder, LADocument, LAFinding, LARisk
from la.parsers import parse_llc1, parse_con29
from la.schemas import LLC1, Con29
from la.risk import run_rules

DATA_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../data"))
os.makedirs(DATA_DIR, exist_ok=True)

def _save_blob(kind: str, content: bytes, order: LAOrder) -> str:
    sha = hashlib.sha256(content).hexdigest()[:10]
    fpath = os.path.join(DATA_DIR, f"{order.id}_{kind}_{sha}.bin")
    with open(fpath, "wb") as f: f.write(content)
    return fpath

def _upsert(db: Session, matter_id: str, key: str, value, ev_file: str, ev_page: int, conf: int = 90):
    db.query(LAFinding).filter(LAFinding.matter_id == matter_id, LAFinding.key == key).delete()
    db.add(LAFinding(matter_id=matter_id, key=key, value=json.dumps(value),
                     evidence_file=ev_file, evidence_page=ev_page, confidence=conf))

def _store_findings_llc1(db: Session, matter: LAMatter, model: LLC1):
    ca = any(c.charge_type == "ConservationArea" and c.present for c in model.charges)
    lb = any(c.charge_type == "ListedBuilding" and c.present for c in model.charges)
    tpo = any(c.charge_type == "TPO" and c.present for c in model.charges)
    a4  = any(c.charge_type == "Article4Direction" and c.present for c in model.charges)
    _upsert(db, matter.id, "llc1.conservation_area.present", bool(ca), "LLC1", 1)
    _upsert(db, matter.id, "llc1.listed_building.present", bool(lb), "LLC1", 1)
    _upsert(db, matter.id, "llc1.tpo.present", bool(tpo), "LLC1", 1)
    _upsert(db, matter.id, "llc1.article4.present", bool(a4), "LLC1", 1)

def _store_findings_con29(db: Session, matter: LAMatter, model: Con29):
    if model.roads_footways is not None:
        _upsert(db, matter.id, "con29.roads_footways.abutting_highway_adopted",
                model.roads_footways.abutting_highway_adopted, "CON29", 1)
        if model.roads_footways.authority:
            _upsert(db, matter.id, "con29.roads_footways.highways_authority",
                    model.roads_footways.authority, "CON29", 1)
    if model.enforcement_notices_present is not None:
        _upsert(db, matter.id, "con29.enforcement_notices_present",
                bool(model.enforcement_notices_present), "CON29", 1)
    if model.contaminated_land_designation is not None:
        _upsert(db, matter.id, "con29.contaminated_land_designation",
                bool(model.contaminated_land_designation), "CON29", 1)
    # NEW: extras
    if model.s106_present is not None:
        _upsert(db, matter.id, "con29.s106_present", bool(model.s106_present), "CON29", 1)
    if model.cil_outstanding is not None:
        _upsert(db, matter.id, "con29.cil_outstanding", bool(model.cil_outstanding), "CON29", 1)
    if model.flood_zone is not None:
        _upsert(db, matter.id, "con29.flood_zone", str(model.flood_zone), "CON29", 1)
    if model.radon_affected is not None:
        _upsert(db, matter.id, "con29.radon_affected", bool(model.radon_affected), "CON29", 1)
    if model.building_regs_completion_present is not None:
        _upsert(db, matter.id, "con29.building_regs_completion_present",
                bool(model.building_regs_completion_present), "CON29", 1)

def parse_and_store(db: Session, matter: LAMatter, docs: List[Tuple[str, bytes]]):
    order = LAOrder(matter_id=matter.id, council_code="MANUAL", types=",".join(k for k,_ in docs), status="READY")
    db.add(order); db.commit(); db.refresh(order)
    for kind, content in docs:
        fpath = _save_blob(kind, content, order)
        db.add(LADocument(order_id=order.id, kind=kind, file_path=fpath, content_type="application/octet-stream"))
    db.commit(); db.refresh(order)

    for doc in order.documents:
        with open(doc.file_path, "rb") as fh: b = fh.read()
        if doc.kind == "LLC1":
            _store_findings_llc1(db, matter, parse_llc1(b))
        elif doc.kind == "CON29":
            _store_findings_con29(db, matter, parse_con29(b))
    db.commit()

    # recompute risks
    db.query(LARisk).filter(LARisk.matter_id == matter.id).delete()
    findings = db.query(LAFinding).filter(LAFinding.matter_id == matter.id).all()
    for r in run_rules(findings):
        r.matter_id = matter.id; db.add(r)
    db.commit()

def generate_client_report(matter: LAMatter, findings: List[LAFinding], risks: List[LARisk], out_path: str) -> str:
    doc = Document()
    doc.add_heading("Local Authority Search Summary", 0)
    doc.add_paragraph(f"Matter: {matter.ref or matter.id}")
    if matter.address: doc.add_paragraph(f"Address: {matter.address}")
    doc.add_paragraph("")

    doc.add_heading("Key Risks", level=1)
    if not risks: doc.add_paragraph("No significant risks identified.")
    else:
        for r in risks:
            p = doc.add_paragraph()
            run = p.add_run(f"[{r.severity}] {r.code}: "); run.bold = True
            p.add_run(r.message or "")

    doc.add_heading("Findings (normalised)", level=1)
    for f in findings:
        try: val = json.loads(f.value) if f.value else None; val_str = json.dumps(val, indent=2)
        except Exception: val_str = f.value or ""
        doc.add_paragraph(f"{f.key}: {val_str}")

    doc.save(out_path); return out_path