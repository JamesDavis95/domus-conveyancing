from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse
from starlette.responses import JSONResponse
from datetime import datetime, timezone

from db import SessionLocal as _Session
from models import Matters, Findings as Finding, Risks as RiskModel, Files as FileModel
from share import verify_token
from llc1_mapper import map_llc1

router = APIRouter()

def _session():
    return _Session()

def _matter(db, mid: int):
    m = db.get(Matters, mid)
    if not m:
        raise HTTPException(404, "matter not found")
    return m

@router.get("/public/passport/{mid}")
def public_passport(mid: int, t: str = Query(..., alias="t")):
    if not verify_token(t, mid):
        raise HTTPException(401, "invalid or expired token")
    db = _session()
    try:
        m = _matter(db, mid)
        if m.status not in ("approved", "published"):
            raise HTTPException(403, "not published")
        prop = m.prop
        findings = db.query(Finding).filter(Finding.matter_id==m.id).all()
        risks = db.query(RiskModel).filter(RiskModel.matter_id==m.id).all()
        return {
            "matter": {
                "id": m.id,
                "ref": m.ref,
                "status": m.status,
                "council": m.council,
                "property": {
                    "uprn": getattr(prop, "uprn", None),
                    "title_no": getattr(prop, "title_no", None),
                    "address_text": getattr(prop, "address_text", None),
                    "postcode": getattr(prop, "postcode", None),
                },
                "findings": [{"type": f.type, "value": f.value, "source_ref": f.source_ref} for f in findings],
                "risks": [{"code": r.code, "severity": r.severity, "explanation": r.explanation} for r in risks],
                "approved_at": m.approved_at.isoformat() if m.approved_at else None,
            }
        }
    finally:
        db.close()

@router.get("/public/llc1/{mid}.docx")
def public_llc1(mid: int, t: str = Query(..., alias="t")):
    if not verify_token(t, mid):
        raise HTTPException(401, "invalid or expired token")
    # re-use logic by calling private builder inline (mirror of llc1_docx)
    from docx import Document
    from tempfile import NamedTemporaryFile
    db = _session()
    try:
        m = _matter(db, mid)
        if m.status not in ("approved", "published"):
            raise HTTPException(403, "not published")
        findings = db.query(Finding).filter(Finding.matter_id==m.id).all()
        risks = db.query(RiskModel).filter(RiskModel.matter_id==m.id).all()
        prop = m.prop
        answers = map_llc1(findings, risks)
        doc = Document()
        doc.add_heading("LLC1 – Official Search of Local Land Charges Register", 0)
        doc.add_paragraph(f"Matter Ref: {m.ref} | Council: {m.council or '-'}")
        if prop:
            doc.add_paragraph(f"Property: {prop.address_text or '-'}  ({prop.postcode or ''})")
            doc.add_paragraph(f"UPRN: {prop.uprn or '-'}  Title: {prop.title_no or '-'}")
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "Charge"; t.rows[0].cells[1].text = "Answer"
        for k in ["CONSERVATION_AREA","LISTED_BUILDING","TREE_PRESERVATION_ORDER","ARTICLE_4_DIRECTION","ENFORCEMENT_NOTICE","FINANCIAL_CHARGES","SMOKE_CONTROL","LIGHTING_CONSENT"]:
            row = t.add_row().cells
            row[0].text = k.replace("_"," ").title()
            row[1].text = answers.get(k, "Unknown")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        return FileResponse(tmp.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="llc1.docx")
    finally:
        db.close()

@router.get("/public/export/{mid}.json")
def public_export_json(mid: int, t: str = Query(..., alias="t")):
    if not verify_token(t, mid):
        raise HTTPException(401, "invalid or expired token")
    db = _session()
    try:
        m = _matter(db, mid)
        if m.status not in ("approved", "published"):
            raise HTTPException(403, "not published")
        # lightweight export: same data shape as /public/passport
        prop = m.prop
        findings = db.query(Finding).filter(Finding.matter_id==m.id).all()
        risks = db.query(RiskModel).filter(RiskModel.matter_id==m.id).all()
        out = {
            "id": m.id, "ref": m.ref, "council": m.council, "status": m.status,
            "property": {
                "uprn": getattr(prop, "uprn", None),
                "title_no": getattr(prop, "title_no", None),
                "address_text": getattr(prop, "address_text", None),
                "postcode": getattr(prop, "postcode", None),
            },
            "findings": [{"type": f.type, "value": f.value, "source_ref": f.source_ref} for f in findings],
            "risks": [{"code": r.code, "severity": r.severity, "explanation": r.explanation} for r in risks],
        }
        return JSONResponse(out)
    finally:
        db.close()
